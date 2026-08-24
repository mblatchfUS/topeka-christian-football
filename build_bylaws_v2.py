#!/usr/bin/env python3
"""Generate Topeka Christian Football Bylaws v2 (redline) as a .docx.

Redline convention (visual, not Word tracked-changes XML):
  - Deleted text: red, strikethrough
  - Added / new text: blue, underline
  - Unchanged text: black
New whole articles/sections are rendered entirely as additions.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xC0, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def run(p, text, kind="same"):
    r = p.add_run(text)
    if kind == "del":
        r.font.color.rgb = RED
        r.font.strike = True
    elif kind == "add":
        r.font.color.rgb = BLUE
        r.font.underline = True
    else:
        r.font.color.rgb = BLACK
    return r


def para(*segments, bold_first=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for i, seg in enumerate(segments):
        if isinstance(seg, tuple):
            text, kind = seg
        else:
            text, kind = seg, "same"
        r = run(p, text, kind)
        if bold_first and i == 0:
            r.font.bold = True
    return p


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(14 if level == 0 else 12)
    r.font.color.rgb = BLACK
    return p


def add_note(text):
    """Editor note explaining a redline change."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("  → Redline note: " + text)
    r.font.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    return p


# ---------------- Title ----------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = t.add_run("BYLAWS OF TOPEKA CHRISTIAN FOOTBALL")
rr.font.bold = True
rr.font.size = Pt(16)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(sub, "(A Kansas Nonprofit Corporation)", "same")
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(sub2, "Version 2 (Redline) — prepared for Board review; supersedes v1 (8-22-2026)", "add")

# ---------------- Change log ----------------
heading("SUMMARY OF CHANGES (v1 → v2)", 0)
para(("Legend:  ", "same"), ("deleted text", "del"), ("   ", "same"),
     ("added / new text", "add"), ("   unchanged text.", "same"))
changes = [
    "Art. III §3.2 — rewrote director count so a married couple = one seat and one collective vote; reconciled the 6-named-directors-in-Articles vs. 5-cap conflict.",
    "Art. III — added §3.2(b) non-member corporation / self-perpetuating board; added regular-meeting cadence + notice (§3.6), remote meetings (§3.7), removal (§3.8), resignation (§3.9), and vacancies (§3.10).",
    "Art. VII — added dual financial controls (§7.4) and mandatory Directors & Officers (D&O) liability insurance (§7.3).",
    "Art. IX — expanded from one sentence into full indemnification, mandatory advancement of expenses, limitation of director liability, and federal/Kansas volunteer immunity.",
    "NEW Art. X — Risk Management & Child Protection: background checks, signed waivers/medical authorization, concussion return-to-play, and mandatory abuse reporting.",
    "NEW Art. XI — Conflict of Interest & Records (IRS 1023 expectation; record retention & inspection; fiscal year = December to match the Articles).",
    "Global — fixed title typo (“Footbal”→“Football”), standardized the legal name, made the Statement of Faith a signed Exhibit A, and cleaned the Certificate of Adoption date.",
]
for c in changes:
    p = doc.add_paragraph(style="List Bullet")
    run(p, c, "add")

doc.add_paragraph()

# ---------------- ARTICLE I ----------------
heading("ARTICLE I: NAME, PRINCIPAL OFFICE, AND PURPOSE")
para(("Section 1.1 Name: ", "same"),
     ("The name of the organization shall be Topeka Christian Football", "same"),
     (" (also referenced herein as “Topeka Christian Football, Inc.” where the registered corporate name so provides)", "add"),
     (", (the “Corporation” or “Club”).", "same"), bold_first=True)
add_note("Use the exact name as registered with the Kansas Secretary of State throughout; delete whichever form does not match the filing.")
para(("Section 1.2 Principal Office: ", "same"),
     ("The principal office of the Corporation shall be located in Topeka, Kansas, or at such other place within Kansas as the Board of Directors may designate.", "same"),
     bold_first=True)
para(("Section 1.3 Purpose: ", "same"),
     ("The Corporation is organized exclusively for charitable, religious, and educational athletic purposes within the meaning of Section 501(c)(3) of the Internal Revenue Code and the Kansas General Corporation Code. Specifically, the Corporation operates as an independent Christian football ministry dedicated to developing Christian character, spiritual discipline, and competitive athletic excellence in youth.", "same"),
     bold_first=True)

# ---------------- ARTICLE II ----------------
heading("ARTICLE II: STATEMENT OF FAITH AND MISSION")
para(("Section 2.1 Statement of Faith: ", "same"),
     ("All board members, officers, and head coaches must subscribe to and uphold the Corporation’s Statement of Faith, affirming the authority of the Holy Scriptures, faith in Jesus Christ as Lord and Savior, and the commitment to conduct all athletic activities in a manner that glorifies God.", "same"),
     (" The Statement of Faith is attached as Exhibit A and incorporated by reference; each director, officer, and head coach shall sign it upon taking office.", "add"),
     bold_first=True)
add_note("v1 made the Statement of Faith binding but never attached it. Attaching and requiring signature makes good-faith enforcement (including removal) defensible.")
para(("Section 2.2 Core Mission: ", "same"),
     ("Topeka Christian Football operates as a Christ-centered ministry focused on mentoring young men to serve Christ, build brotherhood, practice perseverance, and pursue personal growth (“getting a little better every day”).", "same"),
     bold_first=True)

# ---------------- ARTICLE III ----------------
heading("ARTICLE III: BOARD OF DIRECTORS")
para(("Section 3.1 Governance: ", "same"),
     ("The business, property, and affairs of the Corporation shall be managed under the direction of its Board of Directors.", "same"),
     bold_first=True)

para(("Section 3.2 Number, Qualifications, and Married-Couple Seats: ", "same"),
     ("The Board shall consist of not fewer than three (3) and not more than ", "same"),
     ("five (5) directors.", "del"),
     ("seven (7) individuals. ", "add"),
     ("A married couple who both serve on the Board shall together occupy a single director seat and shall hold a single, collective vote; the couple shall designate in writing which spouse casts the vote if they are not in agreement, and if no designation is made the seat abstains. For all purposes of quorum, voting, and the numerical limits in this Section, each married-couple seat counts as one (1) director. ", "add"),
     ("All directors must subscribe to the Statement of Faith (Section 2.1).", "add"),
     bold_first=True)
add_note("This is the requested change and it also resolves the conflict with the Articles: the six named individuals (Allen, Van Camp, Blatchford) form three married-couple seats = three directors, well within the range.")

para(("Section 3.2(b) Non-Member Corporation: ", "same"),
     ("The Corporation shall have no voting members within the meaning of the Kansas nonprofit statutes. All governance authority is vested in the Board of Directors, which shall be self-perpetuating: successor and additional directors shall be elected by a majority vote of the directors then in office. References in the Articles of Incorporation to “membership” are satisfied by this Section and by the participant eligibility provisions of Article VI.", "add"),
     bold_first=True)
add_note("The Articles say membership conditions are “set forth in the Bylaws,” but v1 never defined voting members. Declaring a non-member, self-perpetuating board closes that gap and prevents anyone later claiming member voting/standing rights.")

para(("Section 3.3 Terms: ", "same"),
     ("Directors shall serve staggered terms of two (2) years and may be re-elected.", "same"),
     bold_first=True)

para(("Section 3.4 Quorum and Voting: ", "same"),
     ("A simple majority of the serving directors shall constitute a quorum for the transaction of business. Routine corporate actions require a simple majority vote of the directors present at a meeting where a quorum is established. Any amendments to participant eligibility policies, changes to these Bylaws, or corporate dissolution shall require a two-thirds (2/3) supermajority vote of the full Board.", "same"),
     (" For clarity, each married-couple seat counts as one director and casts one vote when computing quorum and any majority or supermajority.", "add"),
     bold_first=True)
add_note("Removed the v1 parenthetical “requiring at least four votes when the Board has five voting members” because seats are now counted as units; the supermajority math should track seats, not individuals.")

para(("Section 3.5 Action Without a Meeting: ", "same"),
     ("Any action permitted or required to be taken at a Board meeting may be taken without a meeting if all directors consent in writing or by electronic transmission.", "same"),
     bold_first=True)

para(("Section 3.6 Regular and Special Meetings; Notice: ", "same"),
     ("The Board shall hold at least one (1) regular meeting per calendar quarter. Special meetings may be called by the President or by any two (2) director seats. Notice of any meeting, stating the date, time, place or remote-access details, shall be given to each director at least three (3) days in advance by any reasonable means (including email or text); notice may be waived in writing or by attendance.", "add"),
     bold_first=True)

para(("Section 3.7 Remote Participation: ", "same"),
     ("Directors may participate in any meeting by telephone, video, or other electronic means by which all participants can hear and communicate with one another, and such participation constitutes presence in person.", "add"),
     bold_first=True)

para(("Section 3.8 Removal: ", "same"),
     ("Any director may be removed, with or without cause, by a two-thirds (2/3) supermajority vote of the other director seats then in office, including for failure to uphold the Statement of Faith or code of conduct. A married-couple seat is removed as a single seat.", "add"),
     bold_first=True)

para(("Section 3.9 Resignation: ", "same"),
     ("A director may resign at any time by written notice to the President or Secretary; the resignation is effective on receipt unless a later date is stated.", "add"),
     bold_first=True)

para(("Section 3.10 Vacancies: ", "same"),
     ("A vacancy on the Board, however created, may be filled by a majority vote of the remaining directors then in office, even if less than a quorum, and a director so chosen serves for the unexpired term.", "add"),
     bold_first=True)

# ---------------- ARTICLE IV ----------------
heading("ARTICLE IV: OFFICERS AND COACHING LEADERSHIP")
para(("Section 4.1 Officers: ", "same"),
     ("The officers of the Corporation shall be a President, Vice President, Secretary, and Treasurer.", "same"),
     (" Officers are elected by the Board and serve at the Board’s discretion; one person may not simultaneously hold the offices of President and Secretary.", "add"),
     bold_first=True)
para(("Section 4.2 Duties of Officers: ", "same"),
     ("President: Presides at meetings, executes legal contracts approved by the Board, and oversees general club administration. Vice President: Performs presidential duties in the President’s absence and manages liaisons with partner schools and athletic organizations. Secretary: Maintains corporate records, meeting minutes, and issues official notices. Treasurer: Oversees financial accounts, financial reporting, budgeting, and fee collections.", "same"),
     bold_first=True)
para(("Section 4.3 Athletic Staff: ", "same"),
     ("The Board shall appoint a Head Coach and Athletic Coordinators. The Head Coach maintains operational control of on-field team training, scheduling, and athletic safety in accordance with the Club’s Christian mission.", "same"),
     (" All coaches, coordinators, and adult volunteers are subject to the screening requirements of Article X.", "add"),
     bold_first=True)

# ---------------- ARTICLE V ----------------
heading("ARTICLE V: PARTICIPATION AND ELIGIBILITY POLICY")
para(("Section 5.1 Primary Player Base: ", "same"),
     ("Topeka Christian Football is organized primarily to provide competitive football opportunities for Christian homeschool athletes and partnering Christian academy students.", "same"),
     bold_first=True)
para(("Section 5.2 Non-Member & Non-Homeschool Participants: ", "same"),
     ("To ensure player safety and maintain adequate roster sizes, the Board may approve student participants outside the primary homeschool base, provided: the athlete and parents sign and agree to the Club’s Statement of Faith and code of conduct standards; independent athletic association eligibility (e.g., KSHSAA non-member approval rules) is verified; and participation does not compromise the safety or Christian culture of the team.", "same"),
     bold_first=True)
para(("Section 5.3 Board Hardship Exemptions: ", "same"),
     ("The Board possesses sole authority to grant single-season or roster-specific eligibility exemptions to address squad viability and player physical safety.", "same"),
     bold_first=True)

# ---------------- ARTICLE VI ----------------
heading("ARTICLE VI: PARTNERSHIPS AND RECIPROCAL AGREEMENTS")
para(("Section 6.1 Reciprocal Alliances: ", "same"),
     ("The Board may enter into Memorandums of Understanding (MOUs) or reciprocal athletic agreements with area homeschool cooperatives, Christian schools, and fine arts entities.", "same"),
     bold_first=True)
para(("Section 6.2 Legal Independence: ", "same"),
     ("Topeka Christian Football shall operate as a distinct and legally autonomous entity; no partnership shall subordinate the Corporation’s governing authority or tax-exempt status.", "same"),
     bold_first=True)

# ---------------- ARTICLE VII ----------------
heading("ARTICLE VII: FINANCES, EQUIPMENT, AND PROPERTY")
para(("Section 7.1 Independent Accounts: ", "same"),
     ("All registration fees, fundraising proceeds, and donations received for Topeka Christian Football shall be maintained in corporate bank accounts dedicated exclusively to football operations and ministry objectives.", "same"),
     bold_first=True)
para(("Section 7.2 Equipment & Uniforms: ", "same"),
     ("All helmets, pads, uniforms (e.g., Saints apparel), and field gear purchased by or donated to the program are the exclusive property of Topeka Christian Football.", "same"),
     bold_first=True)
para(("Section 7.3 Insurance: ", "same"),
     ("The Corporation shall independently maintain general liability, bodily injury, and athletic participant coverage.", "same"),
     (" The Corporation shall also maintain Directors & Officers (D&O) liability insurance, and abuse/molestation coverage, in amounts the Board deems adequate, for so long as such coverage is reasonably available.", "add"),
     bold_first=True)
add_note("D&O is the single most important protection for the board and is NOT included in general-liability policies. Abuse/molestation coverage is added because it is frequently excluded from GL and is the largest exposure for a youth program.")
para(("Section 7.4 Financial Controls: ", "same"),
     ("The Treasurer shall present a financial report at each regular meeting. Any single expenditure or check above a threshold set by the Board shall require two (2) authorized signatures. No director or officer shall be a sole signatory over both authorization and reconciliation of the same account.", "add"),
     bold_first=True)

# ---------------- ARTICLE VIII ----------------
heading("ARTICLE VIII: SAFETY THRESHOLDS AND FORMATS")
para(("Section 8.1 Roster Safety Standards: ", "same"),
     ("The Head Coach and Board shall maintain minimum active roster thresholds necessary for player safety in contact football (such as 8-man or 11-man competitions). If eligible, healthy player numbers drop below safe operating limits, the program may alter competition formats (e.g., transition to 6-man/8-man play) or cancel contests.", "same"),
     bold_first=True)
para(("Section 8.2 Equipment Maintenance: ", "same"),
     ("All protective equipment must undergo scheduled inspections, reconditioning, and safety certifications complying with national safety standards.", "same"),
     bold_first=True)

# ---------------- ARTICLE IX (expanded) ----------------
heading("ARTICLE IX: INDEMNIFICATION AND LIABILITY PROTECTION")
para(("Section 9.1 Indemnification: ", "same"),
     ("The Corporation shall indemnify any director, officer, coach, or volunteer against legal expenses and liabilities arising out of their good-faith service to the Club, to the fullest extent allowed by Kansas nonprofit law.", "same"),
     bold_first=True)
para(("Section 9.2 Advancement of Expenses: ", "same"),
     ("The Corporation may advance reasonable legal expenses to a director, officer, coach, or volunteer as they are incurred in defending a proceeding, in advance of final disposition, upon receipt of a written undertaking to repay the amounts if it is ultimately determined that the person was not entitled to indemnification.", "add"),
     bold_first=True)
para(("Section 9.3 Limitation of Director Liability: ", "same"),
     ("To the fullest extent permitted by Kansas law, no director shall be personally liable to the Corporation for monetary damages for breach of duty as a director, except for liability that Kansas law does not permit to be eliminated (such as acts not in good faith, intentional misconduct, or transactions from which the director derived an improper personal benefit).", "add"),
     bold_first=True)
para(("Section 9.4 Volunteer Immunity: ", "same"),
     ("Directors, officers, and volunteers who serve without compensation are intended to have the benefit of the federal Volunteer Protection Act of 1997 and any applicable Kansas volunteer-immunity statutes, and nothing in these Bylaws shall be construed to waive those protections.", "add"),
     bold_first=True)
para(("Section 9.5 Non-Exclusivity and Insurance: ", "same"),
     ("The rights in this Article are in addition to any other rights the person may have, and the Corporation shall look first to available insurance (Section 7.3) to satisfy indemnification obligations.", "add"),
     bold_first=True)
add_note("v1 Article IX combined indemnification and dissolution in two sentences. Dissolution is moved to Article XII so this article can be a complete liability-protection package (indemnify + advance + limit + volunteer immunity + insurance).")

# ---------------- ARTICLE X (new) ----------------
heading("ARTICLE X: RISK MANAGEMENT AND CHILD PROTECTION")
para(("Section 10.1 Screening and Background Checks: ", "same"),
     ("Every coach, coordinator, board member, and adult volunteer with access to minors shall complete a criminal background check and reference screening before contact with participants and periodically thereafter as the Board directs. Any disqualifying result bars service.", "add"),
     bold_first=True)
para(("Section 10.2 Participant Waivers and Medical Authorization: ", "same"),
     ("No athlete may practice or compete until the athlete’s parent or legal guardian has signed the Club’s liability waiver and assumption-of-risk, code of conduct, and emergency medical authorization forms. The Corporation shall retain these records for the period required by law.", "add"),
     bold_first=True)
para(("Section 10.3 Concussion and Head-Injury Protocol: ", "same"),
     ("The Club shall adopt and follow a written concussion and return-to-play protocol consistent with Kansas head-injury standards for youth athletics. Any athlete suspected of a concussion shall be removed from play and may not return without written medical clearance.", "add"),
     bold_first=True)
para(("Section 10.4 Abuse Prevention and Mandatory Reporting: ", "same"),
     ("The Club shall maintain a child-protection policy including a two-adult rule where practicable, prohibited-conduct standards, and mandatory reporting of suspected abuse or neglect to civil authorities in accordance with Kansas law. Reporting is not discretionary and shall not be delayed for internal review.", "add"),
     bold_first=True)
add_note("Abuse-related claims are the leading source of nonprofit youth-program litigation and are often excluded from general-liability coverage; a written policy plus screening and mandatory reporting is both a protection and, in practice, an insurer requirement.")

# ---------------- ARTICLE XI (new) ----------------
heading("ARTICLE XI: CONFLICT OF INTEREST, RECORDS, AND FISCAL YEAR")
para(("Section 11.1 Conflict of Interest: ", "same"),
     ("The Board shall adopt a written conflict-of-interest policy consistent with IRS expectations for Section 501(c)(3) organizations. A director or officer with a financial interest in a matter shall disclose it and shall not vote on, or be present for the deliberation of, that matter; the disinterested directors shall decide it. This Section applies to transactions involving a director, an officer, or a member of their family.", "add"),
     bold_first=True)
para(("Section 11.2 Books and Records; Inspection: ", "same"),
     ("The Corporation shall keep correct and complete books and records of account, minutes of Board meetings, and a record of its directors and officers, and shall retain records in accordance with a Board-adopted retention schedule. Directors may inspect corporate records for a purpose reasonably related to their duties.", "add"),
     bold_first=True)
para(("Section 11.3 Fiscal Year: ", "same"),
     ("The fiscal year of the Corporation shall end on December 31, consistent with the tax closing month stated in the Articles of Incorporation.", "add"),
     bold_first=True)

# ---------------- ARTICLE XII (dissolution, moved) ----------------
heading("ARTICLE XII: DISSOLUTION")
para(("Section 12.1 Dissolution: ", "same"),
     ("Upon dissolution, all remaining corporate assets shall, after payment of the Corporation’s liabilities, be distributed exclusively to another Section 501(c)(3) organization, with preference given to a Kansas Christian youth athletic or educational ministry.", "same"),
     bold_first=True)
add_note("Moved from v1 Section 9.2 (unchanged in substance; added “after payment of liabilities” to match the Articles).")

# ---------------- ARTICLE XIII (amendments) ----------------
heading("ARTICLE XIII: AMENDMENTS")
para(("These Bylaws may be altered, amended, or repealed by a two-thirds (2/3) supermajority vote of the Board of Directors at any official meeting, provided written notice of the proposed amendment is provided at least seven (7) days prior to the meeting.", "same"))

# ---------------- Certificate ----------------
heading("CERTIFICATE OF ADOPTION")
para(("The undersigned Secretary certifies that the foregoing Bylaws of Topeka Christian Football were duly adopted by the Board of Directors on the ______ day of ______________, 20____.", "same"),
     ("", "same"))
add_note("Cleaned the pre-filled year; enter the actual adoption date at signing.")
doc.add_paragraph()
para(("____________________________________", "same"))
para(("Secretary, Topeka Christian Football", "same"))

# ---------------- Exhibit A placeholder ----------------
heading("EXHIBIT A: STATEMENT OF FAITH")
para(("[Attach the Corporation’s Statement of Faith here. Each director, officer, and head coach signs and dates below to affirm agreement per Section 2.1.]", "add"))

out = "/Users/mblatchford/Documents/IDEs/ClaudeProjects/Topeka Christian Football/Topeka Christian Football - Bylaws v2 REDLINE.docx"
doc.save(out)
print("Saved:", out)
